import re
from io import BytesIO
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfplumber

def extract_data_from_pdf(file):
    data = {
        "Mã tỉnh": "YBI",
        "Số hóa đơn": "",
        "Mã EVN": "",
        "Mã tháng (yyyyMM)": "",
        "Kỳ": "1",
        "Mã CSHT": "",
        "Ngày đầu kỳ": "",
        "Ngày cuối kỳ": "",
        "Tổng chỉ số": "",
        "Số tiền": "",
        "Thuế VAT": "",
        "Số tiền dự kiến": "",
        "Ghi chú": ""
    }

    with pdfplumber.open(file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"

        # Lấy Số hóa đơn: tìm dòng có "Số (No):"
        m = re.search(r'Số\s*\(No\):\s*(\d+)', text)
        if m:
            data["Số hóa đơn"] = m.group(1).strip()

        # Lấy Mã EVN (Mã khách hàng)
        m = re.search(r'Mã khách hàng \(Customer\'s Code\):\s*([A-Z0-9]+)', text)
        if m:
            data["Mã EVN"] = m.group(1).strip()

        # Lấy Mã CSHT
        m = re.search(r'Tên đơn vị \(Company name\):\s*(.+)', text)
        if m:
            data["Mã CSHT"] = "CSHT_YBI_00014"  # Nếu cố định như mẫu

        # Lấy Mã tháng yyyyMM từ mô tả "Điện tiêu thụ tháng X năm YYYY"
        m = re.search(r'Điện tiêu thụ tháng (\d+) năm (\d{4})', text)
        if m:
            month = int(m.group(1))
            year = int(m.group(2))
            data["Mã tháng (yyyyMM)"] = f"{year}{month:02d}"

        # Lấy Ngày đầu kỳ và Ngày cuối kỳ
        m = re.search(r'Điện tiêu thụ tháng \d+ năm \d{4} từ ngày (\d{2}/\d{2}/\d{4}) đến ngày (\d{2}/\d{2}/\d{4})', text)
        if m:
            data["Ngày đầu kỳ"] = m.group(1)
            data["Ngày cuối kỳ"] = m.group(2)

        # Lấy Tổng chỉ số, Số tiền, Thuế VAT, Số tiền dự kiến (Tổng cộng tiền thanh toán)
        # Tổng chỉ số từ cột Số lượng (Quantity)
        m = re.search(r'Số lượng\s*\(Quantity\)[^\d]*(\d+[\.,]?\d*)', text)
        if m:
            data["Tổng chỉ số"] = m.group(1).replace(",", "")

        # Số tiền (Thành tiền)
        m = re.search(r'(?i)(?:Thành tiền|Cộng tiền hàng).{0,20}(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?)', text)
        if m:
            data["Số tiền"] = m.group(1).replace(",", "").replace(".", "")

        # Thuế VAT
        m = re.search(r'Tiền thuế GTGT \(VAT amount\):\s*(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?)', text)
        if m:
            data["Thuế VAT"] = m.group(1).replace(",", "").replace(".", "")

        # Số tiền dự kiến (Tổng cộng tiền thanh toán)
        m = re.search(r'Tổng cộng tiền thanh toán \(Total payment\):\s*(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?)', text)
        if m:
            data["Số tiền dự kiến"] = m.group(1).replace(",", "").replace(".", "")

    return data

def create_excel(data_list):
    df = pd.DataFrame(data_list)
    # Đảm bảo thứ tự cột
    columns_order = ["Mã tỉnh", "Số hóa đơn", "Mã EVN", "Mã tháng (yyyyMM)", "Kỳ", "Mã CSHT",
                     "Ngày đầu kỳ", "Ngày cuối kỳ", "Tổng chỉ số", "Số tiền", "Thuế VAT", "Số tiền dự kiến", "Ghi chú"]
    df = df[columns_order]

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Hóa đơn')
        worksheet = writer.sheets['Hóa đơn']

        # Tự động giãn cột và định dạng dạng Text
        for idx, col in enumerate(df.columns):
            max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
            worksheet.set_column(idx, idx, max_len)
            worksheet.set_column(idx, idx, None, {'num_format': '@'})  # Định dạng text

    output.seek(0)
    return output.read()

def create_word(data_list):
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    document.styles['Normal'].font.size = Pt(14)

    for data in data_list:
        document.add_paragraph(f"Mã tỉnh: {data.get('Mã tỉnh','')}")
        document.add_paragraph(f"Số hóa đơn: {data.get('Số hóa đơn','')}")
        document.add_paragraph(f"Mã EVN: {data.get('Mã EVN','')}")
        document.add_paragraph(f"Mã tháng (yyyyMM): {data.get('Mã tháng (yyyyMM)','')}")
        document.add_paragraph(f"Kỳ: {data.get('Kỳ','')}")
        document.add_paragraph(f"Mã CSHT: {data.get('Mã CSHT','')}")
        document.add_paragraph(f"Ngày đầu kỳ: {data.get('Ngày đầu kỳ','')}")
        document.add_paragraph(f"Ngày cuối kỳ: {data.get('Ngày cuối kỳ','')}")
        document.add_paragraph(f"Tổng chỉ số: {data.get('Tổng chỉ số','')}")
        document.add_paragraph(f"Số tiền: {data.get('Số tiền','')}")
        document.add_paragraph(f"Thuế VAT: {data.get('Thuế VAT','')}")
        document.add_paragraph(f"Số tiền dự kiến: {data.get('Số tiền dự kiến','')}")
        document.add_paragraph(f"Ghi chú: {data.get('Ghi chú','')}")
        document.add_paragraph("")

    word_buffer = BytesIO()
    document.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer.read()
